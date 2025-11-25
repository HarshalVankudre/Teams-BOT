"""
Document Processor
Handles PDF, JSON, DOCX, Excel, and Images with Vision API
"""
import os
import json
import base64
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from openai import AsyncOpenAI
from .config import config
from .chunker import SemanticChunker
from .vector_store import PineconeStore

# Cache directory for chunks
CHUNK_CACHE_DIR = Path("cache/chunks")


class DocumentProcessor:
    """
    Process various document types and index them into Pinecone.
    - PDF: Extract text and images, describe images with Vision
    - JSON: Each meaningful object becomes its own vector
    - DOCX: Extract text content
    - Excel: Convert to structured data
    - Images: Describe with Vision API
    """

    def __init__(self):
        self.client = AsyncOpenAI(api_key=config.openai_api_key)
        self.chunker = SemanticChunker()
        self.vector_store = PineconeStore()
        # Ensure cache directory exists
        CHUNK_CACHE_DIR.mkdir(parents=True, exist_ok=True)

    def _get_cache_path(self, file_path: Path) -> Path:
        """Get cache file path for a document"""
        # Create hash of file path + modification time for cache key
        stat = file_path.stat()
        cache_key = f"{file_path}_{stat.st_mtime}_{stat.st_size}"
        cache_hash = hashlib.md5(cache_key.encode()).hexdigest()
        return CHUNK_CACHE_DIR / f"{file_path.stem}_{cache_hash}.json"

    def _load_cached_chunks(self, file_path: Path) -> Optional[List[Dict[str, Any]]]:
        """Load chunks from cache if available"""
        cache_path = self._get_cache_path(file_path)
        if cache_path.exists():
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                print(f"Loaded {len(data['chunks'])} chunks from cache")
                return data["chunks"]
            except Exception as e:
                print(f"Cache load error: {e}")
        return None

    def _save_chunks_to_cache(self, file_path: Path, chunks: List[Dict[str, Any]]):
        """Save chunks to cache"""
        cache_path = self._get_cache_path(file_path)
        try:
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump({"source": str(file_path), "chunks": chunks}, f, ensure_ascii=False, indent=2)
            print(f"Saved {len(chunks)} chunks to cache: {cache_path.name}")
        except Exception as e:
            print(f"Cache save error: {e}")

    async def process_and_index(
        self,
        file_path: str,
        additional_context: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process a file and index it into Pinecone.

        Args:
            file_path: Path to the file
            additional_context: Extra context about the document

        Returns:
            Processing statistics
        """
        file_path = Path(file_path)
        if not file_path.exists():
            return {"error": f"File not found: {file_path}"}

        file_ext = file_path.suffix.lower()
        file_name = file_path.name

        print(f"Processing: {file_name} ({file_ext})")

        # Check for cached chunks first
        chunks = self._load_cached_chunks(file_path)

        if chunks is None:
            # No cache - process the file
            print("No cache found, processing file...")

            # Process based on file type
            if file_ext == ".pdf":
                chunks = await self._process_pdf(file_path, additional_context)
            elif file_ext == ".json":
                chunks = await self._process_json(file_path)
            elif file_ext == ".jsonl":
                chunks = await self._process_jsonl(file_path)
            elif file_ext in [".docx", ".doc"]:
                chunks = await self._process_docx(file_path, additional_context)
            elif file_ext in [".xlsx", ".xls"]:
                chunks = await self._process_excel(file_path, additional_context)
            elif file_ext in [".png", ".jpg", ".jpeg", ".gif", ".webp"]:
                chunks = await self._process_image(file_path)
            elif file_ext in [".txt", ".md"]:
                chunks = await self._process_text(file_path, additional_context)
            else:
                return {"error": f"Unsupported file type: {file_ext}"}

            if not chunks:
                return {"error": "No chunks generated from file"}

            # Save chunks to cache BEFORE embedding (in case embedding fails)
            self._save_chunks_to_cache(file_path, chunks)
        else:
            print(f"Using cached chunks for {file_name}")

        # Index chunks into Pinecone
        result = await self.vector_store.upsert_chunks(chunks)

        return {
            "file": file_name,
            "type": file_ext,
            "chunks_created": len(chunks),
            "chunks_indexed": result.get("upserted_count", 0),
            "status": "success"
        }

    async def _process_pdf(
        self,
        file_path: Path,
        additional_context: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Process PDF with text extraction and image description"""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(str(file_path))
            all_chunks = []

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text
                text = page.get_text()

                # Extract and describe images
                image_descriptions = []
                images = page.get_images()

                for img_index, img in enumerate(images):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)

                        # Convert to PNG bytes
                        if pix.n - pix.alpha > 3:  # CMYK
                            pix = fitz.Pixmap(fitz.csRGB, pix)

                        img_bytes = pix.tobytes("png")

                        # Describe with Vision
                        description = await self._describe_image(
                            img_bytes,
                            context=f"Bild auf Seite {page_num + 1} des Dokuments {file_path.name}"
                        )
                        image_descriptions.append(f"[Bild {img_index + 1}]: {description}")

                    except Exception as e:
                        print(f"Error processing image on page {page_num + 1}: {e}")

                # Combine text and image descriptions
                page_content = text
                if image_descriptions:
                    page_content += "\n\n--- Bilder auf dieser Seite ---\n"
                    page_content += "\n".join(image_descriptions)

                # Semantic chunking for this page
                if page_content.strip():
                    page_context = f"Seite {page_num + 1} von {len(doc)}"
                    if additional_context:
                        page_context = f"{additional_context}. {page_context}"

                    page_chunks = await self.chunker.chunk_document(
                        content=page_content,
                        source_file=file_path.name,
                        doc_type="pdf",
                        additional_context=page_context
                    )

                    # Add page number to each chunk
                    for chunk in page_chunks:
                        chunk["source_page"] = page_num + 1
                        chunk["has_images"] = len(images) > 0

                    all_chunks.extend(page_chunks)

            doc.close()
            return all_chunks

        except ImportError:
            print("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return []
        except Exception as e:
            print(f"PDF processing error: {e}")
            return []

    async def _process_json(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process JSON where each meaningful object becomes a vector"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                json_data = json.load(f)

            chunks = await self.chunker.chunk_json_document(
                json_data=json_data,
                source_file=file_path.name
            )

            return chunks

        except Exception as e:
            print(f"JSON processing error: {e}")
            return []

    async def _process_jsonl(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process JSONL file - each line is one record/machine"""
        import asyncio

        try:
            # Read all lines
            with open(file_path, "r", encoding="utf-8") as f:
                lines = [line.strip() for line in f if line.strip()]

            print(f"Found {len(lines)} records in JSONL file")

            # Parse all JSON objects
            records = []
            for i, line in enumerate(lines):
                try:
                    obj = json.loads(line)
                    records.append((i, obj))
                except json.JSONDecodeError as e:
                    print(f"  Skipping invalid JSON on line {i+1}: {e}")

            print(f"Parsed {len(records)} valid records")

            # Process in parallel with semaphore
            semaphore = asyncio.Semaphore(20)
            processed = [0]
            total = len(records)

            async def process_record(idx, obj):
                async with semaphore:
                    chunk = await self.chunker._create_json_chunk(
                        obj=obj,
                        json_path=f"line_{idx}",
                        source_file=file_path.name
                    )
                    processed[0] += 1
                    if processed[0] % 100 == 0 or processed[0] == total:
                        print(f"  Progress: {processed[0]}/{total}")
                    return chunk

            tasks = [process_record(idx, obj) for idx, obj in records]
            chunks = await asyncio.gather(*tasks)

            return list(chunks)

        except Exception as e:
            print(f"JSONL processing error: {e}")
            return []

    async def _process_docx(
        self,
        file_path: Path,
        additional_context: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Process Word document"""
        try:
            from docx import Document

            doc = Document(str(file_path))

            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Also extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                full_text.append("\n[Tabelle]\n" + "\n".join(table_text))

            content = "\n\n".join(full_text)

            chunks = await self.chunker.chunk_document(
                content=content,
                source_file=file_path.name,
                doc_type="docx",
                additional_context=additional_context
            )

            return chunks

        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return []
        except Exception as e:
            print(f"DOCX processing error: {e}")
            return []

    async def _process_excel(
        self,
        file_path: Path,
        additional_context: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Process Excel file - each sheet becomes chunks"""
        try:
            import pandas as pd

            all_chunks = []

            # Read all sheets
            xlsx = pd.ExcelFile(str(file_path))

            for sheet_name in xlsx.sheet_names:
                df = pd.read_excel(xlsx, sheet_name=sheet_name)

                # Convert to structured text
                content = f"## Sheet: {sheet_name}\n\n"
                content += df.to_markdown(index=False)

                sheet_context = f"Excel-Sheet: {sheet_name}"
                if additional_context:
                    sheet_context = f"{additional_context}. {sheet_context}"

                chunks = await self.chunker.chunk_document(
                    content=content,
                    source_file=f"{file_path.name}#{sheet_name}",
                    doc_type="excel",
                    additional_context=sheet_context
                )

                for chunk in chunks:
                    chunk["excel_sheet"] = sheet_name

                all_chunks.extend(chunks)

            return all_chunks

        except ImportError:
            print("pandas/openpyxl not installed. Install with: pip install pandas openpyxl")
            return []
        except Exception as e:
            print(f"Excel processing error: {e}")
            return []

    async def _process_image(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process standalone image with Vision API"""
        try:
            with open(file_path, "rb") as f:
                img_bytes = f.read()

            description = await self._describe_image(
                img_bytes,
                context=f"Eigenständiges Bild: {file_path.name}",
                detailed=True
            )

            # Create a single chunk for the image
            chunk = {
                "id": f"img_{file_path.stem}",
                "content": description,
                "title": file_path.name,
                "summary": description[:200],
                "keywords": ["bild", "image", file_path.stem],
                "category": "image",
                "importance": "medium",
                "source_file": file_path.name,
                "doc_type": "image",
                "chunk_type": "image_description",
                "chunk_index": 0,
                "total_chunks": 1
            }

            return [chunk]

        except Exception as e:
            print(f"Image processing error: {e}")
            return []

    async def _process_text(
        self,
        file_path: Path,
        additional_context: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """Process plain text or markdown file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            chunks = await self.chunker.chunk_document(
                content=content,
                source_file=file_path.name,
                doc_type="text" if file_path.suffix == ".txt" else "markdown",
                additional_context=additional_context
            )

            return chunks

        except Exception as e:
            print(f"Text processing error: {e}")
            return []

    async def _describe_image(
        self,
        image_bytes: bytes,
        context: str = "",
        detailed: bool = False
    ) -> str:
        """Describe an image using GPT-5.1 with reasoning"""
        try:
            # Encode image to base64
            base64_image = base64.b64encode(image_bytes).decode("utf-8")

            prompt = """Beschreibe dieses Bild ausführlich für die Dokumentensuche.
Fokussiere auf:
- Textinhalte im Bild
- Diagramme, Charts, Tabellen
- Wichtige visuelle Elemente
- Zahlen und Daten
Gib eine strukturierte Beschreibung auf Deutsch."""

            if detailed:
                prompt += "\n\nSei besonders detailliert bei Tabellen und Diagrammen."

            # Use chunking model for image description
            request_params = {
                "model": config.chunking_model,
                "input": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": f"{context}\n\n{prompt}"},
                            {
                                "type": "input_image",
                                "image_url": f"data:image/png;base64,{base64_image}",
                                "detail": "high" if detailed else "auto"
                            }
                        ]
                    }
                ],
                "max_output_tokens": 1000
            }
            # Only include reasoning if not "none"
            if config.chunking_reasoning and config.chunking_reasoning != "none":
                request_params["reasoning"] = {"effort": config.chunking_reasoning}

            response = await self.client.responses.create(**request_params)

            return response.output_text

        except Exception as e:
            print(f"Vision API error: {e}")
            return f"[Bildbeschreibung fehlgeschlagen: {str(e)}]"

    async def process_directory(
        self,
        directory_path: str,
        recursive: bool = True,
        file_types: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Process all supported files in a directory.

        Args:
            directory_path: Path to the directory
            recursive: Whether to process subdirectories
            file_types: List of file extensions to process (e.g., ['.pdf', '.docx'])

        Returns:
            Processing statistics
        """
        directory = Path(directory_path)
        if not directory.exists():
            return {"error": f"Directory not found: {directory_path}"}

        supported_types = file_types or [
            ".pdf", ".json", ".jsonl", ".docx", ".doc",
            ".xlsx", ".xls", ".txt", ".md",
            ".png", ".jpg", ".jpeg"
        ]

        results = {
            "total_files": 0,
            "successful": 0,
            "failed": 0,
            "files": []
        }

        # Get files
        if recursive:
            files = [f for f in directory.rglob("*") if f.suffix.lower() in supported_types]
        else:
            files = [f for f in directory.glob("*") if f.suffix.lower() in supported_types]

        results["total_files"] = len(files)

        for file_path in files:
            print(f"\nProcessing: {file_path}")
            result = await self.process_and_index(str(file_path))

            if result.get("status") == "success":
                results["successful"] += 1
            else:
                results["failed"] += 1

            results["files"].append(result)

        return results
